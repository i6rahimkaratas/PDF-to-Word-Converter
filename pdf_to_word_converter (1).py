from pdf2docx import Converter
import os
import fitz
from docx import Document
from docx.shared import Inches
import pdfplumber

def pdf_to_word_pdf2docx(pdf_path, word_path):
    try:
        cv = Converter(pdf_path)
        cv.convert(word_path, start=0, end=None)
        cv.close()
        print(f"Başarıyla dönüştürüldü: {word_path}")
        return True
    except Exception as e:
        print(f"Hata: {e}")
        return False

def pdf_to_word_pymupdf(pdf_path, word_path):
    try:
        doc = fitz.open(pdf_path)
        word_doc = Document()
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text()
            
            if page_num > 0:
                word_doc.add_page_break()
            
            if text.strip():
                word_doc.add_paragraph(text)
            
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                xref = img[0]
                pix = fitz.Pixmap(doc, xref)
                
                if pix.n - pix.alpha < 4:
                    img_data = pix.tobytes("png")
                    img_name = f"temp_img_{page_num}_{img_index}.png"
                    
                    with open(img_name, "wb") as img_file:
                        img_file.write(img_data)
                    
                    word_doc.add_picture(img_name, width=Inches(6))
                    os.remove(img_name)
                
                pix = None
        
        word_doc.save(word_path)
        doc.close()
        
        print(f"Başarıyla dönüştürüldü: {word_path}")
        return True
        
    except Exception as e:
        print(f"Hata: {e}")
        return False

def pdf_to_word_pdfplumber(pdf_path, word_path):
    try:
        word_doc = Document()
        
        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                if page_num > 0:
                    word_doc.add_page_break()
                
                text = page.extract_text()
                if text:
                    word_doc.add_paragraph(text)
                
                tables = page.extract_tables()
                for table in tables:
                    if table:
                        word_table = word_doc.add_table(rows=len(table), cols=len(table[0]))
                        word_table.style = 'Table Grid'
                        
                        for i, row in enumerate(table):
                            for j, cell in enumerate(row):
                                if cell:
                                    word_table.cell(i, j).text = str(cell)
        
        word_doc.save(word_path)
        print(f"Başarıyla dönüştürüldü: {word_path}")
        return True
        
    except Exception as e:
        print(f"Hata: {e}")
        return False

def main():
    pdf_file = "ornek.pdf"
    word_file = "cikti.docx"
    
    print("PDF to Word Dönüştürücü")
    print("=" * 30)
    
    if not os.path.exists(pdf_file):
        print(f"Hata: {pdf_file} dosyası bulunamadı!")
        return
    
    print("Hangi yöntemi kullanmak istiyorsunuz?")
    print("1. pdf2docx (Önerilen - en iyi sonuç)")
    print("2. PyMuPDF + python-docx (Görsel desteği)")
    print("3. pdfplumber + python-docx (Tablo desteği)")
    
    choice = input("Seçiminiz (1-3): ").strip()
    
    if choice == "1":
        success = pdf_to_word_pdf2docx(pdf_file, word_file)
    elif choice == "2":
        success = pdf_to_word_pymupdf(pdf_file, word_file)
    elif choice == "3":
        success = pdf_to_word_pdfplumber(pdf_file, word_file)
    else:
        print("Geçersiz seçim!")
        return
    
    if success:
        print("\nDönüştürme işlemi tamamlandı!")
    else:
        print("\nDönüştürme işlemi başarısız!")

def simple_convert(pdf_path, word_path):
    try:
        cv = Converter(pdf_path)
        cv.convert(word_path)
        cv.close()
        print("Dönüştürme başarılı!")
    except Exception as e:
        print(f"Hata: {e}")

if __name__ == "__main__":
    main()

def batch_convert_folder(input_folder, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    
    pdf_files = [f for f in os.listdir(input_folder) if f.lower().endswith('.pdf')]
    
    for pdf_file in pdf_files:
        pdf_path = os.path.join(input_folder, pdf_file)
        word_file = pdf_file.replace('.pdf', '.docx')
        word_path = os.path.join(output_folder, word_file)
        
        print(f"Dönüştürülüyor: {pdf_file}")
        pdf_to_word_pdf2docx(pdf_path, word_path)
    
    print(f"Toplu dönüştürme tamamlandı! {len(pdf_files)} dosya işlendi.")